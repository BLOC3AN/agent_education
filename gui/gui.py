import streamlit as st
import requests 
from typing import Any, Dict
from src.utils.logger import Logger
logger=Logger(__name__)

import re
import base64

# --- Cấu hình chung ---
# Cấu hình page chỉ chạy một lần
if "page_config_set" not in st.session_state:
    st.set_page_config(page_title="Agent Education", page_icon="🎓")
    st.session_state.page_config_set = True

# Định nghĩa URL của Agent FastAPI của bạn
# Đảm bảo Agent FastAPI đang chạy tại địa chỉ và cổng này
AGENT_API_URL = "http://localhost:2222/agent/request"

# --- Lớp GUI ---
class GUI:
    def __init__(self):
        # Không cần khởi tạo AgentConversation ở đây nữa vì nó là dịch vụ FastAPI riêng
        pass

    def init_gui(self):
        """Khởi tạo giao diện người dùng chính của Streamlit."""
        st.title("🎓 Agent Education - AI Giáo dục")
        st.markdown("*Hệ thống AI hỗ trợ giáo dục và tư vấn học tập*")

        # Khởi tạo session state cho docx_content nếu chưa có
        if "docx_content" not in st.session_state:
            st.session_state.docx_content = {}  # Dictionary để lưu nội dung docx cho mỗi tin nhắn
            st.session_state.docx_filename = {}  # Dictionary để lưu tên file cho mỗi tin nhắn
            st.session_state.custom_filenames = {}  # Dictionary để lưu tên file tùy chỉnh cho mỗi tin nhắn
        
        # Khởi tạo session state cho hiển thị dialog nhập tên file
        if "show_filename_dialog" not in st.session_state:
            st.session_state.show_filename_dialog = False
            st.session_state.dialog_message_index = -1

        # Sidebar để chọn mode và các tùy chọn khác
        with st.sidebar:
            st.header("⚙️ Cài đặt")
            # Nút toggle cho chế độ streaming (chỉ là giao diện, logic thực tế sẽ nhận full response)
            streaming_mode = st.toggle("🔄 Streaming Mode", value=True, help="Bật để hiển thị hiệu ứng 'đang suy nghĩ' và đợi phản hồi đầy đủ. (Lưu ý: API hiện tại không hỗ trợ streaming từng token)")

            if st.button("🗑️ Xóa lịch sử chat"):
                st.session_state.messages = []
                st.session_state.docx_content = {}
                st.session_state.docx_filename = {}
                st.session_state.custom_filenames = {}
                st.session_state.show_filename_dialog = False
                st.session_state.dialog_message_index = -1
                st.rerun() # Làm mới lại ứng dụng để xóa lịch sử

        # Khởi tạo session state cho messages nếu chưa có
        if "messages" not in st.session_state:
            st.session_state.messages = []

        # Hiển thị lịch sử chat
        for i, message in enumerate(st.session_state.messages):
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
                
                # Chỉ hiển thị tùy chọn tải xuống cho tin nhắn của assistant
                if message["role"] == "assistant":
                    # Nếu tin nhắn này chưa có docx content, tạo ngay
                    if i not in st.session_state.docx_content:
                        docx_content, docx_filename = self._convert_md_to_docx(message["content"])
                        if docx_content:
                            st.session_state.docx_content[i] = docx_content
                            st.session_state.docx_filename[i] = docx_filename
                    
                    # Nếu tin nhắn này có docx content, hiển thị nút tải xuống
                    if i in st.session_state.docx_content:
                        with st.expander("📥 Tải xuống tài liệu"):
                            # Lấy tên file tùy chỉnh từ session state hoặc sử dụng tên mặc định
                            default_filename = st.session_state.custom_filenames.get(i, st.session_state.docx_filename.get(i, "tai_lieu.docx"))
                            
                            # Text input cho phép người dùng nhập tên file
                            custom_filename = st.text_input(
                                "Tên file:",
                                value=default_filename,
                                key=f"filename_input_{i}"
                            )
                            
                            # Đảm bảo filename có đuôi .docx
                            if not custom_filename.lower().endswith('.docx'):
                                custom_filename += '.docx'
                            
                            # Lưu tên file tùy chỉnh vào session state
                            st.session_state.custom_filenames[i] = custom_filename
                            
                            # Nút tải xuống
                            st.download_button(
                                label="📥 Tải tài liệu",
                                data=base64.b64decode(st.session_state.docx_content[i]),
                                file_name=custom_filename,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"download_docx_{i}",
                                use_container_width=True,
                            )

        # Ô nhập liệu cho người dùng
        if prompt := st.chat_input("Hãy hỏi gì đó về giáo dục, học tập hoặc sức khỏe..."):
            # Thêm tin nhắn người dùng vào lịch sử chat
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)

            # Xử lý phản hồi từ AI Agent (gọi API)
            with st.chat_message("assistant"):
                if streaming_mode:
                    self._handle_response_via_api_with_spinner(prompt)
                else:
                    self._handle_response_via_api_direct(prompt)

    def _call_agent_api(self, user_input: str):
        """
        Gửi yêu cầu HTTP POST đến Agent FastAPI và trả về phản hồi.
        """
        try:
            payload = {"input": user_input} # Dữ liệu cần gửi trong body của POST request

            # Gửi POST request đến Agent API
            response = requests.post(AGENT_API_URL, json=payload, timeout=60) # Thêm timeout
            response.raise_for_status() # Ném HTTPError cho mã trạng thái lỗi (4xx hoặc 5xx)

            # Phân tích phản hồi JSON từ Agent
            return response.json()

        except requests.exceptions.Timeout:
            return {"error": "API Timeout: Agent không phản hồi kịp thời."}
        except requests.exceptions.ConnectionError:
            return {"error": "Lỗi kết nối: Không thể kết nối tới Agent API. Đảm bảo Agent đang chạy."}
        except requests.exceptions.HTTPError as e:
            return {"error": f"Lỗi HTTP từ Agent: {e.response.status_code} - {e.response.text}"}
        except requests.exceptions.RequestException as e:
            return {"error": f"Lỗi yêu cầu: {e}"}
        except Exception as e:
            return {"error": f"Lỗi không xác định khi gọi API: {e}"}

    def _handle_response_via_api_with_spinner(self, prompt):
        """
        Xử lý phản hồi bằng cách gọi API, hiển thị spinner và update đầy đủ khi có kết quả.
        (Mô phỏng "streaming" nhưng thực tế là chờ full response từ API).
        """
        message_placeholder = st.empty() # Placeholder để cập nhật tin nhắn
        debug_placeholder = st.empty() # Placeholder cho debug/loading info

        message_placeholder.markdown("🤔 Agent đang suy nghĩ... ")
        debug_placeholder.info("Đang chờ phản hồi từ Agent API...")

        try:
            api_result = self._call_agent_api(prompt)
            logger.info(f"Result from agent : {api_result.keys()}")

            if "error" in api_result:
                full_response = f"❌ Lỗi: {api_result['error']}"
                message_placeholder.error(full_response)
                debug_placeholder.empty() # Xóa debug placeholder
            else:
                # Giả sử Agent FastAPI trả về {"response": "..."}
                response_content = api_result.get("response", "Không nhận được phản hồi hợp lệ từ Agent.")
                full_response:Any|Dict = response_content
                if isinstance(full_response, dict) and "output" in full_response:
                    with message_placeholder.container():
                        st.markdown(full_response['output']) # Hiển thị toàn bộ phản hồi
                
                    # Thêm phản hồi vào lịch sử chat
                    st.session_state.messages.append({"role": "assistant", "content": full_response['output']})
                
                    # Tạo docx content sẽ được xử lý trong init_gui
                
                else:
                    message_placeholder.warning(f"⚠️ Agent phản hồi không có định dạng mong đợi: {full_response}")
                    # Thêm phản hồi vào lịch sử chat
                    st.session_state.messages.append({"role": "assistant", "content": str(full_response)})
                
                debug_placeholder.empty() # Xóa debug placeholder sau khi có kết quả

        except Exception as e:
            error_msg = f"❌ Đã xảy ra lỗi không mong muốn trong Streamlit: {str(e)}"
            message_placeholder.error(error_msg)
            debug_placeholder.empty()
            st.session_state.messages.append({"role": "assistant", "content": error_msg})


    def _handle_response_via_api_direct(self, prompt):
        """
        Xử lý phản hồi bằng cách gọi API và hiển thị kết quả trực tiếp (không có hiệu ứng streaming).
        """
        with st.spinner("Đang suy nghĩ..."):
            try:    
                api_result = self._call_agent_api(prompt)
            
                if "error" in api_result:
                    response_content = f"❌ Lỗi: {api_result['error']}"
                    st.error(response_content)
                    # Thêm phản hồi vào lịch sử chat
                    st.session_state.messages.append({"role": "assistant", "content": response_content})
                else:
                    # Giả sử Agent FastAPI trả về {"response": "..."}
                    response_content = api_result.get("response", "Không nhận được phản hồi hợp lệ từ Agent.")
                    
                    # Tạo container để hiển thị phản hồi
                    response_container = st.container()
                    
                    with response_container:
                        if isinstance(response_content, dict) and "output" in response_content:
                            st.markdown(response_content["output"])
                            # Thêm phản hồi vào lịch sử chat
                            st.session_state.messages.append({"role": "assistant", "content": response_content["output"]})
                            
                            # Tạo docx content sẽ được xử lý trong init_gui
                        
                        else:
                            st.markdown(response_content)
                            # Thêm phản hồi vào lịch sử chat
                            st.session_state.messages.append({"role": "assistant", "content": response_content})
                            
                            # Tạo docx content sẽ được xử lý trong init_gui

            except Exception as e:
                error_msg = f"❌ Đã xảy ra lỗi không mong muốn trong Streamlit: {str(e)}"
                st.error(error_msg)
                st.session_state.messages.append({"role": "assistant", "content": error_msg})

    def _check_for_document(self, response_content: str):
        """
        Kiểm tra xem phản hồi có phải là tài liệu có thể tải xuống hay không.
        Nếu có, tạo file DOCX và lưu vào session state.
        """
        # Kiểm tra các pattern phổ biến của tài liệu giáo dục
        document_patterns = [
            r"đề\s+thi",
            r"kiểm\s+tra",
            r"bài\s+kiểm\s+tra",
            r"giáo\s+án",
            r"bài\s+giảng",
            r"bài\s+tập",
            r"tài\s+liệu",
            r"kế\s+hoạch\s+giảng\s+dạy",
            r"giáo\s+trình",
            r"đề\s+cương"
        ]
        
        is_document = any(re.search(pattern, response_content.lower()) for pattern in document_patterns)
        
        if is_document:
            # Tạo tên file mặc định
            default_filename = "tai_lieu_giao_duc.docx"
            
            # Tìm tên file phù hợp từ nội dung
            filename_match = re.search(r"(de_thi|giao_an|bai_tap|tai_lieu)_[\w_]+\.docx", response_content)
            if filename_match:
                filename = filename_match.group(0)
            else:
                # Xác định loại tài liệu
                doc_type = "tai_lieu"
                if re.search(r"đề\s+thi", response_content.lower()):
                    doc_type = "de_thi"
                elif re.search(r"giáo\s+án", response_content.lower()):
                    doc_type = "giao_an"
                elif re.search(r"bài\s+tập", response_content.lower()):
                    doc_type = "bai_tap"
                
                # Tìm môn học
                subject_patterns = {
                    "toán": "toan",
                    "tiếng việt": "tieng_viet",
                    "tiếng anh": "tieng_anh",
                    "văn": "van",
                    "lịch sử": "lich_su",
                    "địa lý": "dia_ly",
                    "khoa học": "khoa_hoc",
                    "vật lý": "vat_ly",
                    "hóa học": "hoa_hoc",
                    "sinh học": "sinh_hoc"
                }
                
                subject = "chung"
                for subj_pattern, subj_code in subject_patterns.items():
                    if re.search(r'\b' + subj_pattern + r'\b', response_content.lower()):
                        subject = subj_code
                        break
                
                # Tìm thông tin lớp
                class_match = re.search(r"lớp\s+(\d+)", response_content.lower())
                class_num = class_match.group(1) if class_match else ""
                
                # Tạo tên file
                if class_num:
                    filename = f"{doc_type}_{subject}_lop{class_num}.docx"
                else:
                    filename = f"{doc_type}_{subject}.docx"
            
            # Chuyển đổi MD sang DOCX
            docx_content, docx_filename = self._convert_md_to_docx(response_content, filename)
            
            if docx_content:
                st.session_state.docx_content = docx_content
                st.session_state.docx_filename = docx_filename
                st.session_state.show_download = True
                # Lưu index của tin nhắn cuối cùng có tài liệu
                st.session_state.last_document_message_index = len(st.session_state.messages) - 1
                
                # Debug log để kiểm tra
                logger.info(f"Document detected! Index: {st.session_state.last_document_message_index}")
                
                return True
        
        return False

    def _convert_md_to_docx(self, message, output_filename=None):
        """
        Convert markdown content to docx file
        Args:
            message: Markdown content to convert
            output_filename: Name of output file (without extension)
        """
        try:
            # Nếu không có output_filename, tạo tên mặc định
            if not output_filename:
                from datetime import datetime
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_filename = f"tai_lieu_{timestamp}.docx"
            elif not output_filename.lower().endswith('.docx'):
                output_filename += '.docx'
                
            # Tạo document từ markdown
            import io
            import base64
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import markdown
            from bs4 import BeautifulSoup
            
            # Chuyển markdown sang HTML với các extension
            html = markdown.markdown(message, extensions=['tables', 'fenced_code', 'nl2br'])
            
            # Tạo document mới
            doc = Document()
            
            # Thiết lập style cho document
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            
            # Thiết lập margin
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Parse HTML
            soup = BeautifulSoup(html, 'html.parser')
            
            # Thêm tiêu đề
            h1_tags = soup.find_all('h1')
            if h1_tags:
                heading = doc.add_heading(h1_tags[0].text, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                heading_format = heading.runs[0].font
                heading_format.color.rgb = RGBColor(0, 0, 128)  # Navy blue
            
            # Xử lý từng phần tử HTML
            for element in soup.find_all(['p', 'h2', 'h3', 'h4', 'ul', 'ol', 'table', 'pre', 'blockquote']):
                if element.name == 'h2':
                    heading = doc.add_heading(element.text, 1)
                    heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)  # Blue
                elif element.name == 'h3':
                    heading = doc.add_heading(element.text, 2)
                    heading.runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Green
                elif element.name == 'h4':
                    heading = doc.add_heading(element.text, 3)
                elif element.name == 'ul':
                    for li in element.find_all('li'):
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(li.text)
                elif element.name == 'ol':
                    for i, li in enumerate(element.find_all('li')):
                        p = doc.add_paragraph(style='List Number')
                        p.add_run(li.text)
                elif element.name == 'table':
                    # Tạo bảng
                    rows = element.find_all('tr')
                    if rows:
                        cols_count = max(len(row.find_all(['td', 'th'])) for row in rows)
                        table = doc.add_table(rows=len(rows), cols=cols_count)
                        table.style = 'Table Grid'
                        
                        # Điền dữ liệu vào bảng
                        for i, row in enumerate(rows):
                            cells = row.find_all(['td', 'th'])
                            for j, cell in enumerate(cells):
                                if j < cols_count:  # Đảm bảo không vượt quá số cột
                                    table.cell(i, j).text = cell.text
                        
                        # Thêm khoảng trống sau bảng
                        doc.add_paragraph()
                elif element.name == 'pre':
                    # Xử lý code block
                    code = element.text
                    p = doc.add_paragraph()
                    code_run = p.add_run(code)
                    code_run.font.name = 'Courier New'
                    code_run.font.size = Pt(10)
                elif element.name == 'blockquote':
                    # Xử lý blockquote
                    p = doc.add_paragraph()
                    p.style = 'Quote'
                    p.add_run(element.text)
                else:
                    # Xử lý paragraph thông thường
                    p = doc.add_paragraph()
                    
                    # Xử lý các phần tử con (bold, italic, links, etc.)
                    for child in element.children:
                        if child.name == 'strong' or child.name == 'b':
                            p.add_run(child.text).bold = True
                        elif child.name == 'em' or child.name == 'i':
                            p.add_run(child.text).italic = True
                        elif child.name == 'a':
                            p.add_run(child.text).underline = True
                        elif child.name == 'code':
                            code_run = p.add_run(child.text)
                            code_run.font.name = 'Courier New'
                        else:
                            # Xử lý text thông thường
                            if hasattr(child, 'text'):
                                p.add_run(child.text)
                            else:
                                p.add_run(str(child))
            
            # Lưu document vào memory
            docx_io = io.BytesIO()
            doc.save(docx_io)
            docx_io.seek(0)
            
            # Encode to base64
            docx_base64 = base64.b64encode(docx_io.read()).decode('utf-8')
            
            return docx_base64, output_filename
            
        except Exception as e:
            logger.error(f"❌ Lỗi khi tạo file DOCX: {str(e)}")
            return None, None

# --- Hàm chạy ứng dụng Streamlit ---
def run_gui():
    gui = GUI()
    gui.init_gui()

# --- Điểm bắt đầu của ứng dụng ---
if __name__ == "__main__":
    run_gui()
